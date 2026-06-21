"""Born-digital Word (.docx) rendering via python-docx.

Determinism: a .docx is a zip, and both the zip member timestamps and the
``docProps/core.xml`` created/modified dates are otherwise wall-clock. We pin the
core-property dates to a fixed datetime and re-pack the archive with normalized
(sorted names, fixed timestamp) members, so the same content -> byte-identical
.docx. This mirrors the ``SOURCE_DATE_EPOCH`` trick used for PDFs.
"""

from __future__ import annotations

import io
from pathlib import Path

from .. import COMPANY_NAME, SYNTHETIC_MARKER
from ._repro import FIXED_DT as _FIXED_DT
from ._repro import normalize_zip as _normalize_docx


def _new_document():
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = Document()
    # Synthetic banner (red, centered) — visible marker on the page.
    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = banner.add_run(SYNTHETIC_MARKER)
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    return doc


def write_contract_docx(ctx: dict, out_path: Path) -> Path:
    """Render the full policy contract to a reproducible .docx."""
    from docx.shared import Pt

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = _new_document()
    doc.add_heading(COMPANY_NAME, level=0)
    doc.add_heading(ctx["doc_title"], level=1)

    intro = doc.add_paragraph()
    intro.add_run(
        f"This policy is a contract between {ctx['holder_name']} (the \"named insured\") and "
        f"{COMPANY_NAME} (\"we\", \"us\"). In return for payment of the premium and subject to all "
        f"terms of this policy, we agree to provide the insurance described herein for "
        f"{ctx['line_label']} policy {ctx['policy_id']}, effective {ctx['effective_date']} through "
        f"{ctx['expiry_date']}."
    )

    for heading, paragraphs in ctx["sections"]:
        doc.add_heading(heading, level=2)
        for para in paragraphs:
            doc.add_paragraph(para)

    closing = doc.add_paragraph()
    closing.add_run(
        "In witness whereof, the Company has caused this policy to be executed. "
        "This document is synthetic and generated for retrieval/benchmark testing; "
        "no real person, business, or policy is described."
    ).italic = True

    return _finalize(doc, out_path, f"{ctx['policy_id']} — {ctx['doc_title']}")


def _finalize(doc, out_path: Path, title: str) -> Path:
    """Pin core-property dates + synthetic subject, then re-pack deterministically."""
    cp = doc.core_properties
    cp.author = COMPANY_NAME
    cp.created = _FIXED_DT
    cp.modified = _FIXED_DT
    cp.subject = SYNTHETIC_MARKER
    cp.title = title

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    bio = io.BytesIO()
    doc.save(bio)
    out_path.write_bytes(_normalize_docx(bio.getvalue()))
    return out_path


def write_sections_docx(kbdoc: dict, out_path: Path) -> Path:
    """Render a generic ``{title, intro, sections}`` document to a reproducible .docx.

    Used by the knowledge base; sections are ``{heading, paragraphs[], bullets[]?}``.
    """
    doc = _new_document()
    doc.add_heading(kbdoc["title"], level=0)
    subtitle = doc.add_paragraph()
    subtitle.add_run(COMPANY_NAME).italic = True
    if kbdoc.get("intro"):
        doc.add_paragraph(kbdoc["intro"])
    for section in kbdoc["sections"]:
        doc.add_heading(section["heading"], level=1)
        for para in section.get("paragraphs", []):
            doc.add_paragraph(para)
        for bullet in section.get("bullets", []):
            doc.add_paragraph(bullet, style="List Bullet")
    return _finalize(doc, out_path, kbdoc["title"])
