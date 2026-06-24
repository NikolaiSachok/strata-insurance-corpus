"""Redaction ground-truth index (#12): deterministic, accurate (text-grounded), and complete.

The PII index is a pure function of the model, but its *accuracy* is what matters for a
redaction benchmark — so we extract the rendered text of every born-digital document and assert
(a) every catalogued text span actually appears (no false ground truth) and (b) every known model
PII value present in a document is catalogued (no missing ground truth).
"""

from __future__ import annotations

import json

import pytest

from generator.content import _eu_date
from generator.model import build_model, index
from generator.run import generate


@pytest.fixture(scope="module")
def corpus(tmp_path_factory):
    out = tmp_path_factory.mktemp("pii_corpus")
    generate(42, out, "sample")
    spans = [json.loads(l) for l in (out / "pii-index.jsonl").read_text().splitlines() if l.strip()]
    manifest = {d["doc_id"]: d for d in json.loads((out / "manifest.json").read_text())["documents"]}
    return out, spans, manifest


def _doc_text(path):
    """Extract text from a born-digital document (PDF/docx/md/csv), or None if unsupported/unavailable."""
    ext = path.suffix.lower()
    if ext == ".pdf":
        import pypdfium2 as pdfium

        doc = pdfium.PdfDocument(str(path))
        return "\n".join(pg.get_textpage().get_text_range() for pg in doc)
    if ext == ".docx":
        import docx

        d = docx.Document(str(path))
        parts = [p.text for p in d.paragraphs]
        parts += [c.text for t in d.tables for r in t.rows for c in r.cells]
        return "\n".join(parts)
    if ext in (".md", ".csv"):
        return path.read_text(encoding="utf-8")
    return None


def _present(value, text, pii_type):
    if str(value) in text:
        return True
    if pii_type == "PERSON_NAME":  # titles/suffixes differ across surfaces — a long token suffices
        return any(w in text for w in str(value).split() if w.isalpha() and len(w) >= 4)
    return False


def test_pii_index_text_grounded(corpus):
    """Every catalogued TEXT span's value must actually appear in the rendered document."""
    out, spans, manifest = corpus
    try:
        import pypdfium2  # noqa: F401
    except Exception as e:  # pragma: no cover
        pytest.skip(f"pypdfium2 unavailable: {e}")

    cache, misses = {}, []
    for s in spans:
        if s["modality"] != "text":
            continue
        path = out / manifest[s["doc_id"]]["path"]
        text = cache.get(path) or cache.setdefault(path, _doc_text(path))
        if text is None:
            continue
        if not _present(s["value"], text, s["pii_type"]):
            misses.append((s["doc_type"], s["pii_type"], s["field"], s["value"]))
    assert not misses, f"catalogued PII not found in rendered text (false ground truth): {misses[:10]}"


def test_pii_index_complete_for_model_values(corpus):
    """Every known model PII value present in a born-digital document must be catalogued."""
    out, spans, manifest = corpus
    try:
        import pypdfium2  # noqa: F401
    except Exception as e:  # pragma: no cover
        pytest.skip(f"pypdfium2 unavailable: {e}")

    m = build_model(42, "sample")
    idx = index(m)
    by_doc: dict[str, set] = {}
    for s in spans:
        by_doc.setdefault(s["doc_id"], set()).add(str(s["value"]))

    def holder_values(h):
        return {h.name, h.email, h.phone, h.street, h.city, h.postcode, h.national_id, _eu_date(h.dob)}

    gaps, cache = [], {}
    for d in manifest.values():
        if d.get("is_scanned") or d.get("is_generated"):
            continue
        path = out / d["path"]
        text = cache.get(path) or cache.setdefault(path, _doc_text(path))
        if text is None:
            continue
        catalogued = by_doc.get(d["doc_id"], set())
        candidates = set()
        for eid in d["entity_ids"]:
            if eid in idx["policyholders"]:
                candidates |= holder_values(idx["policyholders"][eid])
            if eid in idx["agents"]:
                candidates.add(idx["agents"][eid].name)
            if eid in idx["adjusters"]:
                candidates.add(idx["adjusters"][eid].name)
            if eid in idx["policies"]:
                candidates.add(eid)
                v = idx["policies"][eid].vehicle
                if v:
                    candidates.add(v["registration"])
            if eid.startswith("C-"):
                candidates.add(eid)
        for val in candidates:
            if val and str(val) in text and str(val) not in catalogued:
                gaps.append((d["doc_type"], val))
    assert not gaps, f"known PII present in a document but NOT catalogued (missing ground truth): {gaps[:10]}"


def test_pii_index_deterministic(corpus, tmp_path):
    """Same (seed, profile) -> byte-identical pii-index.jsonl."""
    out, _, _ = corpus
    generate(42, tmp_path / "again", "sample")
    assert (out / "pii-index.jsonl").read_bytes() == (tmp_path / "again" / "pii-index.jsonl").read_bytes()


def test_pii_index_structure_and_id_card(corpus):
    """Structural integrity + the ID card exposes its full PII surface."""
    out, spans, manifest = corpus
    valid_mod = {"text", "image_text", "image_region"}
    holder_ids = {h["id"] for h in json.loads((out / "model.json").read_text())["policyholders"]}
    for s in spans:
        assert s["doc_id"] in manifest, f"span references unknown doc {s['doc_id']}"
        assert s["modality"] in valid_mod
        # only the FACE image-region span may have a null value; all redaction text targets are non-empty
        if s["pii_type"] == "FACE":
            assert s["value"] is None and s["modality"] == "image_region"
        else:
            assert s["value"], f"empty PII value in {s}"

    # an ID card carries name + DOB + national id + card no + address + MRZ for its holder
    for hid in holder_ids:
        card = [s for s in spans if s["doc_id"] == f"DOC-{hid}-IDCARD"]
        types = {s["pii_type"] for s in card}
        assert {"PERSON_NAME", "DATE_OF_BIRTH", "NATIONAL_ID", "ID_DOCUMENT_NUMBER", "ADDRESS", "MRZ"} <= types, hid
        # the matching portrait is catalogued as a FACE image-region
        face = [s for s in spans if s["doc_id"] == f"DOC-{hid}-FACE"]
        assert face and face[0]["pii_type"] == "FACE"

    # scanned variants inherit their source's spans as image-text redaction targets
    scanned = [s for s in spans if s["is_scanned"]]
    assert scanned and all(s["modality"] == "image_text" for s in scanned)
